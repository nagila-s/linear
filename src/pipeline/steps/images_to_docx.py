"""
Monta um .docx com imagens extraídas (uma por bloco, com legenda opcional da página).
"""

from __future__ import annotations

import io
import re
import threading
from collections.abc import Callable, Iterable

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Mm, Pt
from PIL import Image

from src.pipeline.steps.pdf_images import ExtractedPdfImage

SUPPORTED_FORMATS = {"PNG", "JPEG", "BMP", "GIF"}
DEFAULT_MAX_WIDTH_IN = 6.0
DEFAULT_MAX_HEIGHT_IN = 8.0

# phase: "iniciando" | "inserindo" | "salvando" | "concluido"
DocxProgressCallback = Callable[[int, int, str], None]


def build_docx_from_images(
    images: Iterable[ExtractedPdfImage],
    *,
    title: str = "Imagens extraídas do PDF",
    include_page_caption: bool = True,
    max_width_in: float = DEFAULT_MAX_WIDTH_IN,
    max_height_in: float = DEFAULT_MAX_HEIGHT_IN,
    on_progress: DocxProgressCallback | None = None,
    progress_every: int = 10,
    save_heartbeat_seconds: float = 8.0,
) -> bytes:
    image_list = images if isinstance(images, list) else list(images)
    total = len(image_list)
    progress_every = max(1, progress_every)

    def report(current: int, phase: str) -> None:
        if on_progress is not None:
            on_progress(current, total, phase)

    report(0, "iniciando")
    doc = Document()
    _configure_document_layout(doc)
    _add_title(doc, _normalize_title(title))

    for index, img in enumerate(image_list, start=1):
        if include_page_caption:
            caption = doc.add_paragraph(f"Página {img.page_number}")
            if caption.runs:
                caption.runs[0].font.name = "Arial"
                caption.runs[0].font.size = Pt(11)
            caption.alignment = WD_ALIGN_PARAGRAPH.LEFT

        pic_data = _normalize_image_bytes(img.image_bytes)
        try:
            width_in, height_in = _fit_image_size_inches(pic_data, max_width_in, max_height_in)
            doc.add_picture(pic_data, width=Inches(width_in), height=Inches(height_in))
        except Exception:
            pic_data.seek(0)
            doc.add_picture(pic_data, width=Inches(max_width_in))

        spacer = doc.add_paragraph()
        spacer.paragraph_format.after = Pt(12)

        if index == 1 or index == total or index % progress_every == 0:
            report(index, "inserindo")

    buffer = io.BytesIO()
    report(total, "salvando")
    _save_with_heartbeat(doc, buffer, interval_seconds=save_heartbeat_seconds, on_tick=on_progress, total=total)
    report(total, "concluido")
    return buffer.getvalue()


def _save_with_heartbeat(
    doc: Document,
    buffer: io.BytesIO,
    *,
    interval_seconds: float,
    on_tick: DocxProgressCallback | None,
    total: int,
) -> None:
    if interval_seconds <= 0 or on_tick is None:
        doc.save(buffer)
        return

    stop = threading.Event()
    tick_count = 0

    def heartbeat() -> None:
        nonlocal tick_count
        while not stop.wait(interval_seconds):
            tick_count += 1
            on_tick(total, total, "salvando")

    thread = threading.Thread(target=heartbeat, daemon=True)
    thread.start()
    try:
        doc.save(buffer)
    finally:
        stop.set()


def _configure_document_layout(doc: Document) -> None:
    section = doc.sections[0]
    section.page_width = Mm(210)
    section.page_height = Mm(297)
    section.top_margin = Mm(12.7)
    section.bottom_margin = Mm(12.7)
    section.left_margin = Mm(12.7)
    section.right_margin = Mm(12.7)
    normal = doc.styles["Normal"]
    normal.font.name = "Arial"
    normal.font.size = Pt(12)


def _add_title(doc: Document, title: str) -> None:
    paragraph = doc.add_paragraph()
    run = paragraph.add_run(title)
    run.font.name = "Arial"
    run.font.size = Pt(16)
    run.bold = True
    paragraph.paragraph_format.after = Pt(20)


def _normalize_title(raw_title: str) -> str:
    title = (raw_title or "").strip()
    title = re.sub(r"\.(docx|doc|pdf)\s*$", "", title, flags=re.IGNORECASE)
    return title.lstrip(" -–—:|_").strip() or "Imagens do PDF"


def _fit_image_size_inches(
    pic_data: io.BytesIO,
    max_width_in: float,
    max_height_in: float,
) -> tuple[float, float]:
    pic_data.seek(0)
    with Image.open(pic_data) as img:
        width_px, height_px = img.size
    pic_data.seek(0)
    if width_px <= 0 or height_px <= 0:
        return max_width_in, max_height_in
    width_in = width_px / 96.0
    height_in = height_px / 96.0
    scale = min(max_width_in / width_in, max_height_in / height_in, 1.0)
    return width_in * scale, height_in * scale


def _normalize_image_bytes(raw: bytes) -> io.BytesIO:
    img = Image.open(io.BytesIO(raw))
    original_format = img.format
    if img.mode in {"RGBA", "LA", "P"} or original_format not in SUPPORTED_FORMATS:
        converted = img.convert("RGB")
        output = io.BytesIO()
        converted.save(output, format="PNG", optimize=True)
        output.seek(0)
        return output
    output = io.BytesIO()
    save_format = "JPEG" if original_format == "JPG" else (original_format or "PNG")
    save_kwargs: dict = {"optimize": True}
    if save_format == "JPEG":
        save_kwargs["quality"] = 88
    img.save(output, format=save_format, **save_kwargs)
    output.seek(0)
    return output
